from volcenginesdkarkruntime import Ark
import os
from docx import Document
import pandas as pd
import docx
from docx.shared import Pt  # 用于设置字号
from docx.shared import RGBColor  # 用于设置字体颜色
from docx.oxml.ns import qn  # 用于设置中文字体


def read_info_from_excel(file_path):
    df = pd.read_excel(file_path)
    datas = []
    for index, row in df.iterrows():
        info = {
            "import_file_name": row["要导入的文件名"],
            "prompt": row["提示语"],
            "export_template": row["导出的模版"],
            "export_file_name": row["导出的文件名"],
        }
        datas.append(info)
    return datas


def read_word_file(file_path):
    # 打开Word文档
    doc = Document(file_path)
    # 初始化一个空字符串来存储文本内容
    full_text = []

    # 遍历文档中的每个段落
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 将所有段落的文本合并为一个字符串，段落之间用换行符分隔
    return "\n".join(full_text)


def combine_word_files(old_file_name, new_string, new_file_name):
    # 读取旧文件内容
    doc = docx.Document(old_file_name)
    paragraphs = [p for p in doc.paragraphs]

    # 创建新文档
    new_doc = docx.Document()

    # 写入旧文件的段落和格式
    for paragraph in paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb

    # 添加新的无格式字符串
    new_paragraph = new_doc.add_paragraph(new_string)

    # 设置字体为 15 号宋体
    style = new_doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(15)
    font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
    r = style.element.rPr.rFonts
    r.set(qn("w:eastAsia"), "宋体")

    # 应用样式到新段落
    for run in new_paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 保存新文件
    new_doc.save("output/" + new_file_name)


def get_answer(input_text, input_prompt):
    client = Ark(
        base_url="https://ark.cn-beijing.volces.com/api/v3",
        ak="AKLTY2E0NWIyYmM0OWE0NGUwOWFlNGE5MjRjYmVmZjc5YmU",
        sk="T0dGaU16azRNV1ptWWpZME5HUXpNMkV5TVdRd1pXSTJPRGxsTkRnME9HVQ==",
    )
    # client.model_config["protected_namespaces"] = ()
    # Non-streaming:
    print("----- standard request -----")
    completion = client.chat.completions.create(
        model="ep-20240805101445-qmrdx",
        messages=[
            {
                "role": "system",
                "content": "你是豆包，是由字节跳动开发的 AI 人工智能助手",
            },
            {"role": "user", "content": input_text + "\n" + input_prompt},
        ],
    )
    return completion.choices[0].message.content


def main():
    datas = read_info_from_excel("./配置文件.xlsx")
    client = Ark(
        base_url="https://ark.cn-beijing.volces.com/api/v3",
        ak="AKLTY2E0NWIyYmM0OWE0NGUwOWFlNGE5MjRjYmVmZjc5YmU",
        sk="T0dGaU16azRNV1ptWWpZME5HUXpNMkV5TVdRd1pXSTJPRGxsTkRnME9HVQ==",
    )
    for data in datas:
        input_text = read_word_file(data["import_file_name"])
        input_prompt = read_word_file(data["prompt"])
        old_file_name = data["export_template"]
        # Non-streaming:
        print("----- standard request -----")
        completion = client.chat.completions.create(
            model="ep-20240805111046-s6znj",
            messages=[
                {
                    "role": "system",
                    "content": "你是豆包，是由字节跳动开发的 AI 人工智能助手",
                },
                {"role": "user", "content": input_text + "\n" + input_prompt},
            ],
        )
        new_string = completion.choices[0].message.content
        new_file_name = data["export_file_name"]
        combine_word_files(
            old_file_name=old_file_name,
            new_string=new_string,
            new_file_name=new_file_name,
        )


if __name__ == "__main__":
    main()
